"""
Document Processing Module

Handles parsing and extraction from PDF and DOCX documents.
Includes metadata extraction and citation management.

Author: Fractal Agent System
License: MIT
"""

import logging
import mimetypes
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict, Any, List, Union

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    logging.warning("PyPDF2 not installed. Install with: pip install PyPDF2")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logging.warning("pdfplumber not installed. Install with: pip install pdfplumber")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed. Install with: pip install python-docx")


logger = logging.getLogger(__name__)


@dataclass
class DocumentMetadata:
    """Metadata extracted from a document."""

    title: Optional[str] = None
    author: Optional[str] = None
    subject: Optional[str] = None
    keywords: Optional[str] = None
    creator: Optional[str] = None
    producer: Optional[str] = None
    creation_date: Optional[datetime] = None
    modification_date: Optional[datetime] = None
    page_count: int = 0
    word_count: int = 0
    file_size: int = 0
    file_type: str = "unknown"
    custom_properties: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary representation."""
        return {
            "title": self.title,
            "author": self.author,
            "subject": self.subject,
            "keywords": self.keywords,
            "creator": self.creator,
            "producer": self.producer,
            "creation_date": self.creation_date.isoformat() if self.creation_date else None,
            "modification_date": self.modification_date.isoformat() if self.modification_date else None,
            "page_count": self.page_count,
            "word_count": self.word_count,
            "file_size": self.file_size,
            "file_type": self.file_type,
            "custom_properties": self.custom_properties
        }


@dataclass
class ParsedDocument:
    """Represents a parsed document with content and metadata."""

    file_path: Path
    content: str
    metadata: DocumentMetadata
    parsed_at: datetime = field(default_factory=datetime.now)
    sections: List[Dict[str, Any]] = field(default_factory=list)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    images: List[Dict[str, Any]] = field(default_factory=list)

    def to_citation(self, style: str = "apa") -> str:
        """
        Generate formatted citation.

        Args:
            style: Citation style ("apa", "mla", "chicago")

        Returns:
            Formatted citation string
        """
        author = self.metadata.author or "Unknown Author"
        title = self.metadata.title or self.file_path.stem
        year = self.metadata.creation_date.year if self.metadata.creation_date else "n.d."

        if style == "apa":
            return f"{author}. ({year}). {title}."
        elif style == "mla":
            return f"{author}. \"{title}.\" {year}."
        elif style == "chicago":
            return f"{author}. {title}. {year}."
        else:
            return f"{author}. {title}. {year}."

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary representation."""
        return {
            "file_path": str(self.file_path),
            "content": self.content,
            "metadata": self.metadata.to_dict(),
            "parsed_at": self.parsed_at.isoformat(),
            "sections": self.sections,
            "tables": self.tables,
            "images": self.images,
            "citation": self.to_citation()
        }

    def get_summary(self, max_length: int = 500) -> str:
        """
        Get document summary.

        Args:
            max_length: Maximum summary length

        Returns:
            Summary string
        """
        content = self.content.strip()
        if len(content) <= max_length:
            return content
        return content[:max_length] + "..."


class PDFParser:
    """PDF document parser with multiple backend support."""

    def __init__(self, use_pdfplumber: bool = True):
        """
        Initialize PDF parser.

        Args:
            use_pdfplumber: Use pdfplumber if available (better text extraction)
        """
        self.use_pdfplumber = use_pdfplumber and PDFPLUMBER_AVAILABLE

        if not PYPDF2_AVAILABLE and not PDFPLUMBER_AVAILABLE:
            raise RuntimeError("No PDF parsing library available")

    def parse(self, file_path: Union[str, Path]) -> ParsedDocument:
        """
        Parse PDF document.

        Args:
            file_path: Path to PDF file

        Returns:
            ParsedDocument object

        Raises:
            FileNotFoundError: If file doesn't exist
            RuntimeError: If parsing fails
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if self.use_pdfplumber:
            return self._parse_with_pdfplumber(file_path)
        else:
            return self._parse_with_pypdf2(file_path)

    def _parse_with_pypdf2(self, file_path: Path) -> ParsedDocument:
        """Parse PDF using PyPDF2."""
        if not PYPDF2_AVAILABLE:
            raise RuntimeError("PyPDF2 not available")

        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)

                metadata = self._extract_metadata_pypdf2(reader, file_path)

                content_parts = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        content_parts.append(text)

                content = "\n\n".join(content_parts)
                metadata.page_count = len(reader.pages)
                metadata.word_count = len(content.split())

                return ParsedDocument(
                    file_path=file_path,
                    content=content,
                    metadata=metadata
                )

        except Exception as e:
            logger.error(f"PyPDF2 parsing failed for {file_path}: {e}")
            raise RuntimeError(f"PDF parsing failed: {e}") from e

    def _parse_with_pdfplumber(self, file_path: Path) -> ParsedDocument:
        """Parse PDF using pdfplumber."""
        if not PDFPLUMBER_AVAILABLE:
            raise RuntimeError("pdfplumber not available")

        try:
            with pdfplumber.open(file_path) as pdf:
                metadata = self._extract_metadata_pdfplumber(pdf, file_path)

                content_parts = []
                tables = []

                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        content_parts.append(text)

                    page_tables = page.extract_tables()
                    for table_num, table in enumerate(page_tables, 1):
                        tables.append({
                            "page": page_num,
                            "table_number": table_num,
                            "data": table
                        })

                content = "\n\n".join(content_parts)
                metadata.page_count = len(pdf.pages)
                metadata.word_count = len(content.split())

                return ParsedDocument(
                    file_path=file_path,
                    content=content,
                    metadata=metadata,
                    tables=tables
                )

        except Exception as e:
            logger.error(f"pdfplumber parsing failed for {file_path}: {e}")
            raise RuntimeError(f"PDF parsing failed: {e}") from e

    def _extract_metadata_pypdf2(
        self,
        reader: 'PyPDF2.PdfReader',
        file_path: Path
    ) -> DocumentMetadata:
        """Extract metadata using PyPDF2."""
        metadata = DocumentMetadata(file_type="pdf")

        if reader.metadata:
            metadata.title = reader.metadata.get('/Title')
            metadata.author = reader.metadata.get('/Author')
            metadata.subject = reader.metadata.get('/Subject')
            metadata.creator = reader.metadata.get('/Creator')
            metadata.producer = reader.metadata.get('/Producer')

            creation_date = reader.metadata.get('/CreationDate')
            if creation_date:
                try:
                    metadata.creation_date = self._parse_pdf_date(creation_date)
                except Exception:
                    pass

            mod_date = reader.metadata.get('/ModDate')
            if mod_date:
                try:
                    metadata.modification_date = self._parse_pdf_date(mod_date)
                except Exception:
                    pass

        metadata.file_size = file_path.stat().st_size

        return metadata

    def _extract_metadata_pdfplumber(
        self,
        pdf: 'pdfplumber.PDF',
        file_path: Path
    ) -> DocumentMetadata:
        """Extract metadata using pdfplumber."""
        metadata = DocumentMetadata(file_type="pdf")

        if pdf.metadata:
            metadata.title = pdf.metadata.get('Title')
            metadata.author = pdf.metadata.get('Author')
            metadata.subject = pdf.metadata.get('Subject')
            metadata.creator = pdf.metadata.get('Creator')
            metadata.producer = pdf.metadata.get('Producer')

            creation_date = pdf.metadata.get('CreationDate')
            if creation_date:
                try:
                    metadata.creation_date = self._parse_pdf_date(creation_date)
                except Exception:
                    pass

            mod_date = pdf.metadata.get('ModDate')
            if mod_date:
                try:
                    metadata.modification_date = self._parse_pdf_date(mod_date)
                except Exception:
                    pass

        metadata.file_size = file_path.stat().st_size

        return metadata

    @staticmethod
    def _parse_pdf_date(date_str: str) -> datetime:
        """Parse PDF date format (D:YYYYMMDDHHmmSS)."""
        if date_str.startswith('D:'):
            date_str = date_str[2:]

        date_str = date_str.split('+')[0].split('-')[0].split('Z')[0]

        formats = [
            '%Y%m%d%H%M%S',
            '%Y%m%d%H%M',
            '%Y%m%d'
        ]

        for fmt in formats:
            try:
                return datetime.strptime(date_str[:len(fmt)], fmt)
            except ValueError:
                continue

        raise ValueError(f"Unable to parse PDF date: {date_str}")


class DOCXParser:
    """DOCX document parser."""

    def __init__(self):
        """Initialize DOCX parser."""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx not available")

    def parse(self, file_path: Union[str, Path]) -> ParsedDocument:
        """
        Parse DOCX document.

        Args:
            file_path: Path to DOCX file

        Returns:
            ParsedDocument object

        Raises:
            FileNotFoundError: If file doesn't exist
            RuntimeError: If parsing fails
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            doc = DocxDocument(file_path)

            metadata = self._extract_metadata(doc, file_path)

            content_parts = []
            sections = []
            tables = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    content_parts.append(text)

                    if para.style.name.startswith('Heading'):
                        sections.append({
                            "level": para.style.name,
                            "text": text
                        })

            for table_num, table in enumerate(doc.tables, 1):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)

                tables.append({
                    "table_number": table_num,
                    "data": table_data
                })

            content = "\n\n".join(content_parts)
            metadata.word_count = len(content.split())

            return ParsedDocument(
                file_path=file_path,
                content=content,
                metadata=metadata,
                sections=sections,
                tables=tables
            )

        except Exception as e:
            logger.error(f"DOCX parsing failed for {file_path}: {e}")
            raise RuntimeError(f"DOCX parsing failed: {e}") from e

    def _extract_metadata(
        self,
        doc: 'DocxDocument',
        file_path: Path
    ) -> DocumentMetadata:
        """Extract metadata from DOCX document."""
        metadata = DocumentMetadata(file_type="docx")

        core_props = doc.core_properties

        metadata.title = core_props.title
        metadata.author = core_props.author
        metadata.subject = core_props.subject
        metadata.keywords = core_props.keywords
        metadata.creator = core_props.author
        metadata.creation_date = core_props.created
        metadata.modification_date = core_props.modified

        metadata.file_size = file_path.stat().st_size

        return metadata


class DocumentParser:
    """Unified document parser supporting multiple formats."""

    def __init__(self):
        """Initialize document parser."""
        self.pdf_parser = PDFParser() if (PYPDF2_AVAILABLE or PDFPLUMBER_AVAILABLE) else None
        self.docx_parser = DOCXParser() if DOCX_AVAILABLE else None

    def parse(self, file_path: Union[str, Path]) -> ParsedDocument:
        """
        Parse document based on file type.

        Args:
            file_path: Path to document file

        Returns:
            ParsedDocument object

        Raises:
            ValueError: If file type not supported
            FileNotFoundError: If file doesn't exist
            RuntimeError: If parsing fails
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_type = self._detect_file_type(file_path)

        if file_type == "pdf":
            if not self.pdf_parser:
                raise RuntimeError("PDF parsing not available")
            return self.pdf_parser.parse(file_path)

        elif file_type == "docx":
            if not self.docx_parser:
                raise RuntimeError("DOCX parsing not available")
            return self.docx_parser.parse(file_path)

        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    @staticmethod
    def _detect_file_type(file_path: Path) -> str:
        """
        Detect file type from extension and MIME type.

        Args:
            file_path: Path to file

        Returns:
            File type string
        """
        suffix = file_path.suffix.lower()

        if suffix == '.pdf':
            return 'pdf'
        elif suffix in ['.docx', '.doc']:
            return 'docx'

        mime_type, _ = mimetypes.guess_type(str(file_path))

        if mime_type:
            if 'pdf' in mime_type:
                return 'pdf'
            elif 'word' in mime_type or 'document' in mime_type:
                return 'docx'

        return 'unknown'

    def parse_multiple(
        self,
        file_paths: List[Union[str, Path]]
    ) -> List[ParsedDocument]:
        """
        Parse multiple documents.

        Args:
            file_paths: List of file paths

        Returns:
            List of ParsedDocument objects
        """
        results = []

        for file_path in file_paths:
            try:
                doc = self.parse(file_path)
                results.append(doc)
            except Exception as e:
                logger.error(f"Failed to parse {file_path}: {e}")
                continue

        return results


def parse_document(file_path: Union[str, Path]) -> ParsedDocument:
    """
    Convenience function for quick document parsing.

    Args:
        file_path: Path to document file

    Returns:
        ParsedDocument object
    """
    parser = DocumentParser()
    return parser.parse(file_path)


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)

    import sys

    if len(sys.argv) < 2:
        print("Usage: python document_parser.py <file_path>")
        sys.exit(1)

    file_path = sys.argv[1]

    try:
        doc = parse_document(file_path)

        print(f"File: {doc.file_path}")
        print(f"Type: {doc.metadata.file_type}")
        print(f"Title: {doc.metadata.title}")
        print(f"Author: {doc.metadata.author}")
        print(f"Pages: {doc.metadata.page_count}")
        print(f"Words: {doc.metadata.word_count}")
        print(f"\nCitation: {doc.to_citation()}")
        print(f"\nSummary:\n{doc.get_summary(200)}")

    except Exception as e:
        logger.error(f"Parsing failed: {e}")
        sys.exit(1)
